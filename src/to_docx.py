from docx import Document
import logging

class ToDocX(object):
	"""Convert the transcription to a DocX file and write it out

	Options 

	output_file = 'path to output file' # if blank, defaults to transcription['title'] + 'docx'
	"""
	_options = {}
	def __init__(self, **kwargs):
		self._options.update(kwargs)

	def __call__(self, transcription):
		"""Create a DocX file from the transcription object"""
		logging.info('Outputting to Word file...')
		document = Document()
		document.add_heading(transcription['title'], 0)
		for paragraph in transcription['paragraphs']:
			document.add_paragraph(paragraph['content'])

		filename = transcription['title'] + '.docx'

		if 'output_file' in self._options:
			filename = self._options['output_file']

		document.save(filename)
		return transcription